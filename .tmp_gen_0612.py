#!/usr/bin/env python3
"""Generate tailored CV + cover letter (md + docx) for 2026-06-12 daily run.
Companies: Accenture (ICT Cadet), TikTok-Frontend (Frontend Engineer Intern, TikTok LIVE),
SafetyCulture (GTM Engineer). Reuses the locked blue 2E74B5 / navy Arial styling.
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/sessions/epic-elegant-allen/mnt/job-application/_daily/2026-06-12"
DATE = "12 June 2026"
FONT="Arial"; NAVY=RGBColor(0x1F,0x38,0x64); ACCENT=RGBColor(0x2E,0x74,0xB5)
BODY=RGBColor(0x33,0x33,0x33); META=RGBColor(0x70,0x70,0x70)

CONTACT = "Sydney, NSW, Australia | quochuy.dinh@student.uts.edu.au | 0410 525 293 | [GitHub](https://github.com/Huydinh1205) | [LinkedIn](https://linkedin.com/in/quochuy-dinh) | [Portfolio](https://cv-website-lemon.vercel.app/)"

# ---------- shared CV section text ----------
EDUCATION = """## EDUCATION

**Bachelor of Artificial Intelligence** | University of Technology Sydney (UTS) | Sydney, Australia | Jul 2025 to Present
- Focus on machine learning, deep learning, and applied AI systems
- Database Fundamentals: High Distinction (100/100), full marks across relational algebra, joins, aggregation, and query optimisation

**Bachelor of Computer Science** | Ho Chi Minh City University of Technology (HCMUT) | Aug 2023 to Jun 2025
- GPA: 3.8 / 4.0
- Data Structures and Algorithms (C++): A grade
"""

def experience(order):
    bullets = {
        "backend": "- Contributed backend server modules (TCP, OBS WebSocket, and REST API layers) for a game-research project, enabling real-time messaging and programmatic control across clients",
        "pipelines": "- Built normalisation, cleaning, and feature-extraction pipelines for downstream modelling across large multi-source datasets",
        "ocr": "- Cleaned and analysed noisy multi-source OCR datasets, ran iterative experiments, and translated findings into reproducible, version-controlled workflows",
        "vlm": "- Currently researching Vision-Language Models and generative modelling, designing and running multimodal experiments",
    }
    head = "## EXPERIENCE\n\n**Research Assistant** | University of Technology Sydney | Sydney, NSW | Nov 2025 to Present\n"
    return head + "\n".join(bullets[k] for k in order) + "\n"

# project blocks: key -> markdown
PROJECTS = {
"ecom": """**E-Commerce Web Application** | React, Express.js, Node.js, PostgreSQL, JWT
- Built a full-stack e-commerce platform independently, covering product listing, search and filter, shopping cart, and checkout with order management
- Implemented JWT-based authentication (register, login, protected routes) with bcrypt password hashing
- Designed and normalised a relational PostgreSQL schema (up to 3NF) and optimised queries for CRUD operations
""",
"coffee": """**AI-Powered Coffee Shop Manager Dashboard** | Python, LangChain, Redis, Docker, AWS (planned)
- Designed and built an end-to-end AI briefing system delivering automated morning reports covering revenue trends and operational insights, removing the need for manual data review
- Developed a single LangChain pipeline integrating four AI capabilities: natural language Q&A, revenue summarisation, inventory forecasting, and scheduled auto-generated daily reports
- Implemented a Redis caching layer to serve repeated queries without redundant recomputation, and containerised the system with Docker for consistent deployment
""",
"curric": """**CurricuLLM, Parent-Teacher AI Platform** | Cambridge EduX Hackathon 2026, Team of 4 | TypeScript, React, FastAPI, PostgreSQL
- Built an AI pipeline that auto-generates personalised curriculum progress reports for parents from teacher input, removing manual report writing
- Developed a parent-facing AI chat interface for natural language queries on student progress, served via FastAPI, with real-time multilingual translation
- Contributed full-stack across the React frontend and the FastAPI/PostgreSQL backend (REST API design, database schema)
""",
"transformer": """**Transformer Chatbot from Scratch** | PyTorch, Lightning, FastAPI, Hugging Face Spaces
- Built a decoder-only transformer entirely from scratch, including a custom tokenizer, causal masking, and temperature sampling, without relying on high-level model libraries
- Served the model through a FastAPI backend and deployed it on Hugging Face Spaces, powering a live chatbot widget on a personal portfolio site
- Tuned batching, caching, and sampling on the inference path to keep the deployed widget responsive under real usage
""",
"alpr": """**Smart Parking and Intelligent Traffic Management (ALPR)** | ML Team Lead | PyTorch, Ultralytics YOLO, PARSEQ, ByteTrack, Streamlit
- Leading the machine learning work on a full automatic licence-plate recognition pipeline targeting international publication
- Combined YOLOv11-OBB plate detection, a PARSEQ transformer OCR head, ByteTrack temporal voting, and Zero-DCE++ low-light enhancement
- Benchmarked robustness across CCPD2020 subsets (blur, tilt, night) and built a Streamlit dashboard for real-time violation detection
""",
"frontend": """**Frontend Projects Portfolio** | React, JavaScript, TypeScript, HTML, CSS
- Built a React Tic-Tac-Toe game with move history and undo (deployed on Netlify), managing game state and time-travel through the component tree
- Built a React Music Player and a React Weather App consuming a live weather API, with async fetch and conditional rendering of loading and error states
- Created pure HTML/CSS/JS UI clones focused on responsive layout, cross-browser behaviour, and reusable component structure
""",
"game": """**Game Automation and Streaming Integration Tool** | Python, Flask, FastAPI, TCP, OBS WebSocket, REST API
- Implemented a TCP layer for real-time client-server communication across a modular backend
- Integrated the OBS WebSocket protocol to react to live stream events in real time
- Exposed keyboard and mouse triggers via REST API endpoints, enabling programmatic control from external clients
""",
"som": """**Set-of-Mark Visual Prompting for Tactical Soccer Understanding** | Research, target DICTA 2026 | PyTorch, Gemini, Qwen2.5-VL, YOLO, OpenCV
- Developed a training-free recipe overlaying detector marks (numbered player boxes, team colours, ball markers) onto broadcast frames so a Vision-Language Model can reason about possession, space, passing lanes, and pressing
- Built a Tactical-QA benchmark and ran ablations against an oracle ceiling derived from SoccerNet game-state reconstruction
- Integrated multiple VLM backends behind a common interface for systematic comparison
""",
}

DATA_SECTION = """## DATA AND ANALYSIS PROJECTS

**Game Telemetry Data Analysis** | Python, Pandas, NumPy, Matplotlib, Seaborn
- Ran exploratory data analysis on competitive-game telemetry: cleaning, feature engineering, and statistical analysis of player behaviour and match outcomes, surfacing actionable performance and outlier insights

**OCR Preprocessing Pipeline** | Python, OpenCV, NumPy, PaddleOCR
- Built an image preprocessing pipeline (adaptive thresholding, denoising, deskewing, contrast enhancement) and benchmarked downstream OCR accuracy across configurations

**Housing Price Data Analysis** | Python, Pandas, scikit-learn, Matplotlib
- Performed EDA, missing-value handling, and statistical modelling to identify the strongest predictors of house price, using hypothesis testing, correlation analysis, and regression diagnostics
"""

def projects_section(order):
    return "## SELECTED PROJECTS\n\n" + "\n".join(PROJECTS[k] for k in order)

def cv_markdown(role, summary, proj_order, skills, include_data=True):
    parts = ['---', 'title: "QUOC HUY DINH (RYAN)"', '---', '',
             f'# {role}', '', CONTACT, '',
             '## PROFESSIONAL SUMMARY', '', summary, '',
             EDUCATION.strip(), '',
             experience(proj_order.get("exp", ["backend","pipelines","ocr","vlm"])).strip(), '',
             projects_section(proj_order["projects"]).strip(), '']
    if include_data:
        parts += [DATA_SECTION.strip(), '']
    parts += [skills.strip(), '']
    return "\n".join(parts)

# ===================== CV docx builder (from build_cv.py) =====================
def set_font(run, size, color=BODY, bold=False, italic=False):
    run.font.name=FONT; run.font.size=Pt(size); run.font.bold=bold; run.font.italic=italic; run.font.color.rgb=color
    rpr=run._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),FONT)
def p_spacing(p,before=0,after=0,line=1.06):
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line
def bottom_border(p,color="2E74B5",sz=6,space=6):
    ppr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz)); b.set(qn('w:space'),str(space)); b.set(qn('w:color'),color)
    pb.append(b); ppr.append(pb)
def add_hyperlink(p,url,text,size=9.5):
    r_id=p.part.relate_to(url,"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)
    hl=OxmlElement('w:hyperlink'); hl.set(qn('r:id'),r_id); nr=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    rf=OxmlElement('w:rFonts'); [rf.set(qn(a),FONT) for a in ('w:ascii','w:hAnsi','w:cs')]; rPr.append(rf)
    szc=OxmlElement('w:sz'); szc.set(qn('w:val'),str(int(size*2))); rPr.append(szc)
    col=OxmlElement('w:color'); col.set(qn('w:val'),'2E74B5'); rPr.append(col)
    u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u); nr.append(rPr)
    t=OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text=text; nr.append(t); hl.append(nr); p._p.append(hl)
def add_plain(p,text,size=9.5,color=META):
    r=p.add_run(text); set_font(r,size,color)
LINK_RE=re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
def split_bold(text):
    out=[]; i=0
    for m in re.finditer(r'\*\*(.+?)\*\*',text):
        if m.start()>i: out.append((text[i:m.start()],False))
        out.append((m.group(1),True)); i=m.end()
    if i<len(text): out.append((text[i:],False))
    return out or [(text,False)]
def render_contact(doc,line):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_spacing(p,2,7)
    for i,seg in enumerate(s.strip() for s in line.split('|')):
        if i: add_plain(p,'   |   ',9.5,META)
        m=LINK_RE.match(seg)
        if m: add_hyperlink(p,m.group(2),m.group(1))
        else: add_plain(p,seg,9.5,BODY)
    bottom_border(p,"2E74B5",6,8)
def render_bold_line(doc,line):
    m=re.match(r'\*\*(.+?)\*\*(.*)',line); label=m.group(1).strip(); rest=m.group(2).strip()
    p=doc.add_paragraph()
    if label.endswith(':'):
        p_spacing(p,1,2); r=p.add_run(label+' '); set_font(r,9.7,BODY,bold=True); add_plain(p,rest,9.7,BODY)
    else:
        p_spacing(p,7,1); p.paragraph_format.keep_with_next=True
        r=p.add_run(label); set_font(r,10.3,NAVY,bold=True)
        if rest: add_plain(p,'   '+rest.lstrip('|').strip(),9.3,META)
def render_bullet(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p_spacing(p,0,2.2,1.05); p.paragraph_format.left_indent=Cm(0.6)
    for seg,b in split_bold(text):
        r=p.add_run(seg); set_font(r,9.7,BODY,bold=b)
def build_cv_docx(md, out):
    txt=md; name="QUOC HUY DINH (RYAN)"
    mm=re.search(r'title:\s*"?([^"\n]+)"?',txt)
    if mm: name=mm.group(1).strip()
    txt=re.sub(r'^---\n.*?\n---\n','',txt,count=1,flags=re.S)
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Cm(0.9); sec.bottom_margin=Cm(0.9); sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)
    st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(9.7)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_spacing(p,0,1)
    r=p.add_run(name); set_font(r,21,NAVY,bold=True)
    rpr=r._element.get_or_add_rPr(); sp=OxmlElement('w:spacing'); sp.set(qn('w:val'),'20'); rpr.append(sp)
    role_done=False; contact_done=False
    for ln in txt.splitlines():
        s=ln.rstrip()
        if not s.strip(): continue
        if s.startswith('# ') and not role_done:
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_spacing(p,0,3)
            r=p.add_run(s[2:].strip().upper()); set_font(r,12.5,ACCENT,bold=True)
            rpr=r._element.get_or_add_rPr(); sp=OxmlElement('w:spacing'); sp.set(qn('w:val'),'30'); rpr.append(sp)
            role_done=True; continue
        if (not contact_done) and ('@' in s and '|' in s):
            render_contact(doc,s); contact_done=True; continue
        if s.startswith('## '):
            p=doc.add_paragraph(); p_spacing(p,8,3)
            r=p.add_run(s[3:].strip().upper()); set_font(r,11.5,ACCENT,bold=True)
            rpr=r._element.get_or_add_rPr(); sp=OxmlElement('w:spacing'); sp.set(qn('w:val'),'20'); rpr.append(sp)
            bottom_border(p,"2E74B5",6,3); continue
        if s.strip()=='---': continue
        if s.startswith('- '): render_bullet(doc,s[2:].strip()); continue
        if s.startswith('**'): render_bold_line(doc,s); continue
        p=doc.add_paragraph(); p_spacing(p,1,3,1.1)
        for seg,b in split_bold(s):
            r=p.add_run(seg); set_font(r,9.8,BODY,bold=b)
    os.makedirs(os.path.dirname(out),exist_ok=True); doc.save(out); print("CV  ->",out)

# ===================== Cover letter builder =====================
def sf(run,size,color=BODY,bold=False):
    run.font.name=FONT; run.font.size=Pt(size); run.font.bold=bold; run.font.color.rgb=color
    rpr=run._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),FONT)
def cover_header(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_spacing(p,0,1)
    r=p.add_run("QUOC HUY DINH (RYAN)"); sf(r,21,NAVY,bold=True)
    rpr=r._element.get_or_add_rPr(); s=OxmlElement('w:spacing'); s.set(qn('w:val'),'20'); rpr.append(s)
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_spacing(c,1,8)
    parts=["Sydney, NSW, Australia","quochuy.dinh@student.uts.edu.au","0410 525 293",
           ("GitHub","https://github.com/Huydinh1205"),("LinkedIn","https://linkedin.com/in/quochuy-dinh"),
           ("Portfolio","https://cv-website-lemon.vercel.app/")]
    for i,seg in enumerate(parts):
        if i: rr=c.add_run("   |   "); sf(rr,9.5,META)
        if isinstance(seg,tuple): add_hyperlink(c,seg[1],seg[0])
        else: rr=c.add_run(seg); sf(rr,9.5,BODY)
    bottom_border(c,"2E74B5",6,8)
def build_cover(folder,recipient,salutation,body):
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Cm(1.0); sec.bottom_margin=Cm(1.0); sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)
    doc.styles['Normal'].font.name=FONT; doc.styles['Normal'].font.size=Pt(10.3)
    cover_header(doc)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p_spacing(p,4,8); r=p.add_run(DATE); sf(r,10,META)
    for i,line in enumerate(recipient):
        p=doc.add_paragraph(); p_spacing(p,0,0,1.15); r=p.add_run(line); sf(r,10.3,BODY,bold=(i==0))
    p=doc.add_paragraph(); p_spacing(p,8,8); r=p.add_run(salutation); sf(r,10.3,BODY)
    for b in body:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p_spacing(p,0,8,1.2); r=p.add_run(b); sf(r,10.3,BODY)
    p=doc.add_paragraph(); p_spacing(p,8,0); r=p.add_run("Sincerely,"); sf(r,10.3,BODY)
    p=doc.add_paragraph(); p_spacing(p,2,0); r=p.add_run("Quoc Huy Dinh (Ryan)"); sf(r,10.3,NAVY,bold=True)
    out=f"{BASE}/{folder}/CV/Huy-CoverLetter-{folder}-2026-06-12.docx"
    os.makedirs(os.path.dirname(out),exist_ok=True); doc.save(out)
    md=[f"# Cover Letter — {folder}","",DATE,""]+recipient+["",salutation,""]+body+["","Sincerely,","Quoc Huy Dinh (Ryan)"]
    open(out[:-5]+".md","w",encoding="utf-8").write("\n".join(md)+"\n"); print("COV ->",out)

def write_cv(folder, md):
    out=f"{BASE}/{folder}/CV/Huy-CV-{folder}-2026-06-12.md"
    os.makedirs(os.path.dirname(out),exist_ok=True)
    open(out,"w",encoding="utf-8").write(md+"\n")
    build_cv_docx(md, out[:-3]+".docx")

def write_jd(folder, text):
    out=f"{BASE}/{folder}/JD/{text['file']}"
    os.makedirs(os.path.dirname(out),exist_ok=True)
    open(out,"w",encoding="utf-8").write(text['body']+"\n"); print("JD  ->",out)

# ============================ ACCENTURE ============================
acc_summary = ("Artificial Intelligence and Computer Science student at UTS who learns fast and delivers, with hands-on experience across the software delivery lifecycle: building and testing full-stack applications, configuring and deploying systems with Docker, debugging through logs, and documenting reproducible workflows. Comfortable supporting delivery teams, working in Agile, and collaborating with developers, analysts, and non-technical stakeholders. Genuinely interested in software development and testing, cloud and infrastructure, and data and analytics, and eager to grow these skills on real client projects as an ICT Cadet at Accenture.")
acc_skills = """## SKILLS

**Programming Languages:** Python, C++, JavaScript, TypeScript, SQL, HTML/CSS
**Delivery and Practices:** Software Development Lifecycle (SDLC), Agile, testing and troubleshooting, debugging via logs, technical documentation, configuration and deployment, version control
**Frameworks and Libraries:** React, Node.js, Express.js, Flask, FastAPI, Pandas, NumPy
**AI / Machine Learning:** PyTorch, scikit-learn, LangChain, LLMs, NLP, data cleaning, feature extraction
**Databases:** PostgreSQL, Redis
**Tools and Platforms:** Docker, Git, Jupyter, VS Code, Hugging Face Spaces, AWS (foundational)
**Protocols and Integration:** REST APIs, JWT Authentication, TCP/IP, OBS WebSocket"""
acc_md = cv_markdown("ICT Cadet", acc_summary,
    {"exp":["backend","pipelines","ocr","vlm"], "projects":["ecom","coffee","curric","alpr","frontend"]},
    acc_skills)

# ============================ TIKTOK FRONTEND ============================
tk_summary = ("Artificial Intelligence and Computer Science student and front-end focused engineer who builds responsive, component-based user interfaces in React with JavaScript and TypeScript. Hands-on experience with state management, consuming REST and live APIs, cross-browser behaviour, and front-end performance, plus deploying working web apps and interactive widgets to production. Comfortable owning the full UI from design to shipped code and collaborating across regions and teams. Excited to join TikTok LIVE to build high-quality, high-performance interfaces that reach users at global scale.")
tk_skills = """## SKILLS

**Programming Languages:** JavaScript, TypeScript, HTML/CSS, Python, C++
**Frontend:** React, component-based development, state management, responsive design, cross-browser compatibility, front-end performance optimisation, REST and live API integration
**Frameworks and Libraries:** React, Node.js, Express.js, FastAPI, Flask
**Databases:** PostgreSQL, Redis
**Tools and Platforms:** Git, Docker, VS Code, Netlify, Hugging Face Spaces, Jupyter
**AI / Machine Learning:** PyTorch, LangChain, LLMs, NLP (applied in deployed AI-powered front-end widgets)
**Protocols and Integration:** REST APIs, JWT Authentication, WebSocket"""
tk_md = cv_markdown("Frontend Engineer Intern", tk_summary,
    {"exp":["backend","vlm","pipelines","ocr"], "projects":["frontend","ecom","curric","transformer","alpr"]},
    tk_skills)

# ============================ SAFETYCULTURE GTM ENGINEER ============================
sc_summary = ("Artificial Intelligence and Computer Science student and builder who ships practical, AI-powered automations that remove manual work. Hands-on experience designing LLM workflows with LangChain (natural-language Q&A, automated reporting, prompt design, and tool use), building lightweight JavaScript and full-stack applications, and integrating systems through REST APIs and relational databases. Strong bias for action: prototypes quickly, deploys with Docker, and iterates from real usage. Comfortable working across technical and non-technical boundaries, and eager to learn CRM tooling such as Salesforce and cloud data warehouses such as Amazon Redshift to embed automation directly into go-to-market teams.")
sc_skills = """## SKILLS

**AI / Automation:** LangChain, Large Language Models (LLMs), prompt engineering and tool use, AI workflow design, NLP, scheduled automated reporting, PyTorch, scikit-learn
**Programming Languages:** JavaScript, TypeScript, Python, C++, SQL, HTML/CSS
**Integration and Data:** REST APIs, relational databases (PostgreSQL, 3NF), SQL querying, JWT Authentication; familiar with cloud data-warehouse concepts (Redshift), open to Salesforce
**Frameworks and Libraries:** FastAPI, Flask, Express.js, React, Node.js, Pandas, NumPy
**Databases and Caching:** PostgreSQL, Redis
**Tools and Platforms:** Docker, Git, VS Code, Hugging Face Spaces, AWS (foundational)"""
sc_md = cv_markdown("Go-To-Market (GTM) Engineer", sc_summary,
    {"exp":["backend","pipelines","ocr","vlm"], "projects":["coffee","curric","transformer","ecom","frontend"]},
    sc_skills)

write_cv("Accenture", acc_md)
write_cv("TikTok-Frontend", tk_md)
write_cv("SafetyCulture", sc_md)

# ---------------- Cover letters ----------------
build_cover("Accenture",
    ["Technology Recruitment Team","Accenture Australia","Level 4/55 Clarence St, Sydney, NSW 2000"],
    "Dear Recruitment Team,",
    ["I am writing to apply for the ICT Cadet role in Accenture's Technology practice in Sydney. The chance to support delivery teams on real client projects, build foundational skills across systems, applications, and infrastructure, and learn Agile delivery methodologies through structured, hands-on work is exactly the kind of start I am looking for as an Artificial Intelligence and Computer Science student.",
     "The role asks for a basic understanding of the software development lifecycle, strong analytical and problem-solving skills, and the ability to learn quickly in a fast-paced environment, and I have built toward all three. I designed and shipped a full-stack e-commerce application end to end, including the React front end, a Node and Express REST API, JWT authentication, and a normalised PostgreSQL schema, then tested and debugged it through to a working product. I built an AI reporting dashboard and containerised it with Docker for consistent deployment, and across my research assistant work I document reproducible, version-controlled workflows and troubleshoot issues through logs. My interests map closely to the areas you mention: software development and testing, cloud and infrastructure, and data and analytics.",
     "I enjoy supporting a team, contributing to documentation and knowledge sharing, and translating requirements into concrete technical tasks, and I am eager to do that the right way under Accenture's delivery methodologies. I would welcome the opportunity to grow as an ICT Cadet and contribute from day one. Thank you for considering my application."])

build_cover("TikTok-Frontend",
    ["APAC Early Careers Team","TikTok","180 George Street, Sydney, NSW 2000"],
    "Dear Hiring Team,",
    ["I am writing to apply for the Frontend Engineer Intern role with TikTok LIVE (2027 Start). Developing user interfaces for products used at the scale of TikTok LIVE, and creating the ultimate user experience through high-quality design and coding, is precisely the kind of front-end work I want to do as an Artificial Intelligence and Computer Science student.",
     "The role asks for strong front-end programming skills, experience optimising front-end performance and handling browser compatibility, an understanding of component development, and familiarity with at least one MV* framework, and React is my core framework. I built a portfolio of React applications, including a Tic-Tac-Toe game with full move history and undo deployed on Netlify, a music player, and a weather app that consumes a live API with async fetching and clean loading and error states. I also built the complete React front end for a full-stack e-commerce platform and a parent-facing AI chat interface in React and TypeScript for an education platform, and I care about responsive layout, reusable components, and keeping deployed interfaces fast and reliable.",
     "I would be excited to help the TikTok LIVE Foundation and Revenue teams ship interfaces that reach users globally, and to learn from your engineers about building front-end systems at scale. Thank you for your time and consideration; I would welcome the chance to discuss how I can contribute."])

build_cover("SafetyCulture",
    ["Marketing Technology / GTM Engineering Team","SafetyCulture","Sydney, NSW"],
    "Dear Hiring Team,",
    ["I am writing to apply for the GTM Engineer role at SafetyCulture. The mission of this role, embedding technical and AI capability directly into the go-to-market team to remove manual toil and build tools that people actually use every day, is exactly the kind of practical, high-ownership building I love doing.",
     "The role asks for someone who can ship production-ready, AI-powered automations using modern LLM APIs, build lightweight JavaScript applications and integrations, and work with data from a cloud warehouse, with a strong bias for action. I have built in this spirit: an end-to-end AI system with a single LangChain pipeline that delivers natural-language Q&A, revenue summarisation, forecasting, and scheduled automated daily reports, replacing manual review, with a Redis cache and Docker deployment. I built an AI platform that auto-generates personalised reports and answers natural-language questions for non-technical users, and I work comfortably in JavaScript and TypeScript across full-stack projects with REST API integrations and normalised SQL databases. I prototype fast, put work in front of real users, and iterate.",
     "I want to be transparent that I have not yet worked hands-on with Salesforce or Amazon Redshift, but my SQL and integration experience transfers directly and I learn new tooling quickly. I am genuinely excited by SafetyCulture's builder culture and its focus on real AI adoption, and I would welcome the chance to discuss how I can contribute. Thank you for considering my application."])

# ---------------- JD analyses ----------------
write_jd("Accenture", {"file":"Accenture-ICTCadet-JD-Analysis-2026-06-12.md","body":"""# Accenture — ICT Cadet — JD Analysis (12 June 2026)

Source (verified live 12 Jun 2026): https://builtinsydney.au/job/ict-cadet/9714288
Posted: ~2 hours before capture (12 Jun 2026). Location: In-office, Level 4/55 Clarence St, Sydney NSW 2000. Level: Entry.

## Role
Early-career ICT Cadet in Accenture's Technology practice. Support delivery teams implementing/maintaining technology solutions; system analysis, testing (UAT, SIT), troubleshooting, documentation, configuration, deployment, defect tracking; learn Agile; collaborate across developers, analysts, project managers.

## Must-haves
- Currently studying or recently completed a degree in IT, Computer Science, Engineering or related field.
- Basic understanding of IT concepts (SDLC, systems, or infrastructure).
- Strong analytical / problem-solving skills; fast learner; good communication and teamwork.
- Interest in software dev/testing, cloud & infrastructure, data & analytics, or business analysis.

## Eligibility (international student)
No citizenship/PR requirement stated in the posting. Standard Australian work-rights checks likely apply at offer stage. Flag: confirm whether the cadetship requires unrestricted AU work rights before investing heavily.

## Huy fit
Strong. Full-stack build/test/deploy experience (E-Commerce: React + Node/Express REST API + JWT + 3NF PostgreSQL), Docker deployment, log-based debugging, reproducible documentation in RA role, and genuine interest spanning software/testing, cloud, and data & analytics. Cadet is entry-level with structured learning, so year-of-study is not a barrier.

## Apply
Built In "Apply" forwards to Accenture's ATS. Confirm work-rights question on application.
"""})

write_jd("TikTok-Frontend", {"file":"TikTok-FrontendEngineerIntern-JD-Analysis-2026-06-12.md","body":"""# TikTok — Frontend Engineer Intern (TikTok LIVE), 2027 Start (BS/MS) — JD Analysis (12 June 2026)

Source (verified live 12 Jun 2026): https://au.gradconnection.com/employers/tiktok/jobs/tiktok-frontend-engineer-intern-tiktok-live-2027-start-bsms-2/
Closes: 30 Jun 2026, 8:59 pm. Internship period: Dec 2026 - Feb 2027 (onboard by 7 Dec 2026). Location: Sydney (on site). ACCEPTS INTERNATIONAL: Yes.

## Role
Front-end engineer intern on TikTok LIVE (Foundation + Revenue teams). Develop UIs for PC/mobile, build engineering/infra tooling, deliver high-quality user experience for the LIVE ecosystem.

## Must-haves
- Currently pursuing an Undergraduate or Master's in Software Development, CS, Computer Engineering, or related technical discipline. (Undergrad eligible.)
- Solid front-end programming skills; experience solving browser-compatibility issues and optimising front-end performance.
- Understanding of front-end engineering and component development; familiarity with at least one MV* framework.
- Preferred: multi-end (Native/PC/Server) dev, large-scale/open-source project experience, tech blogging.

## Eligibility (international student)
Explicitly "Accepts International". Apply limit: max 2 TikTok roles globally — coordinate with the TikTok Backend (LIVE) application from 11 Jun.

## Huy fit
Strong front-end match: React is his core framework; portfolio of deployed React apps (Tic-Tac-Toe w/ history+undo on Netlify, music player, weather app w/ live API), full React front end for e-commerce, React+TypeScript chat UI for CurricuLLM. Gap vs preferred: large-scale/open-source experience — frame research + hackathon teamwork as collaboration evidence.

## Apply
Via GradConnection "Apply". Note 2-role TikTok cap.
"""})

write_jd("SafetyCulture", {"file":"SafetyCulture-GTMEngineer-JD-Analysis-2026-06-12.md","body":"""# SafetyCulture — GTM Engineer — JD Analysis (12 June 2026)

Source (verified live 12 Jun 2026): https://builtinsydney.au/job/gtm-engineer/8937114
Reposted ~19 days before capture. Location: Hybrid, Sydney NSW. Level: tagged Entry (content reads mid-level).

## Role
Embed technical + AI capability into the Go-To-Market (sales/CS) team. Identify workflow friction, build production AI-powered automations and JavaScript tooling, own integrations between GTM tools (Salesforce, Redshift), drive AI adoption.

## Must-haves
- Proven ability to build automation/integration across GTM stacks, with hands-on Salesforce (custom objects, flows, APIs, data management). [GAP for Huy]
- JavaScript development for lightweight apps, custom integrations, browser-based tooling. [MATCH]
- Querying Amazon Redshift or comparable cloud warehouse. [PARTIAL — SQL transfers, no Redshift]
- Design/ship AI workflows with modern LLM APIs (esp. Anthropic Claude): prompt design, tool use, production deployment. [STRONG MATCH]
- Enough sales/CS process understanding to scope solutions independently. [GAP]

## Eligibility (international student)
No citizenship/PR requirement stated. JD explicitly encourages applying without meeting every requirement.

## Huy fit
Stretch. Strong on the AI-automation core (LangChain pipeline with NL Q&A + automated reporting; LLM tool use; Redis/Docker) and JavaScript/full-stack. Clear gaps: no Salesforce/CRM, no Redshift, limited commercial sales-process exposure. Worth applying given the "apply anyway" framing and the strong AI-builder narrative; be transparent about Salesforce/Redshift as learn-on-the-job.

## Apply
Built In "Apply" forwards to SafetyCulture careers. Lead with AI-automation + JS; address Salesforce/Redshift gap honestly.
"""})

print("DONE")
